"""Word document processor."""

import logging
from pathlib import Path
from typing import Dict, Any

from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


class WordProcessor:
    """Process Word (.docx) documents using python-docx."""

    def __init__(self):
        self.supported_formats = [".doc", ".docx"]

    async def process(self, file_path: str) -> Dict[str, Any]:
        """Extract basic text and metadata from a Word document."""
        try:
            path = Path(file_path)
            if not path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")

            doc = DocxDocument(str(path))
            text = "\n".join(p.text for p in doc.paragraphs)
            word_count = len(text.split())

            metadata = {
                "paragraphs": len(doc.paragraphs)
            }

            return {
                "content_type": "word",
                "metadata": metadata,
                "page_count": None,
                "word_count": word_count,
                "text": text,
                "pages": [],
                "images": [],
                "tables": [],
                "structure": {},
                "brand_elements": {},
            }
        except Exception as e:
            logger.error(f"Error processing Word document {file_path}: {e}")
            raise
